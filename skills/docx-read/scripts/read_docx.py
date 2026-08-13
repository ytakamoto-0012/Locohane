"""docxファイルを読み込み、段落・表・文書プロパティを JSON で出力する。

docx-tools スキルの実行スクリプト（progressive disclosure 第3段階）。
run_script から
    python read_docx.py <docx_path> [--offset N] [--limit N]
の形で呼ばれる。

read_file.py の offset/limit と同じ考え方を段落単位に適用し、
1回の呼び出しで返す段落数の上限を設けて巨大な出力を避ける
（表は全件返す。読み込みのみで書き込みは行わないため python-docx の
「既存ファイルへの変更」という性質は関係ない）。

.doc（レガシーのバイナリ形式）は python-docx が非対応のため扱えない。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from _common import setup_utf8_stdio, summarize_result, write_json_result


def _length_cm(value) -> float | None:
    """python-docxのLength（EMU）をcm単位のfloatへ変換する（None時はNoneのまま）。"""
    return round(value.cm, 2) if value is not None else None


def main() -> int:
    setup_utf8_stdio()
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path")
    parser.add_argument("--offset", type=int, default=0)
    parser.add_argument("--limit", type=int, default=300)
    args = parser.parse_args()

    path = Path(args.docx_path)
    if not path.exists():
        print(f"ファイルが見つかりません: {args.docx_path}", file=sys.stderr)
        return 1
    if path.is_dir():
        print(f"指定パスはディレクトリです（ファイル専用）: {args.docx_path}", file=sys.stderr)
        return 1
    if path.suffix.lower() == ".doc":
        print(
            "拡張子が .doc（レガシーのバイナリ形式）です。このスキルは .docx のみ対応しています。"
            "Microsoft Wordで開き「名前を付けて保存」で .docx 形式に変換してから再度お試しください。",
            file=sys.stderr,
        )
        return 1
    if path.suffix.lower() != ".docx":
        print(f"拡張子が .docx ではありません: {args.docx_path}", file=sys.stderr)
        return 1

    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        print(f".docxとして読み込めませんでした（壊れているか非対応形式の可能性）: {args.docx_path}", file=sys.stderr)
        return 1
    except OSError as e:
        print(f"ファイル読み込みに失敗しました: {e}", file=sys.stderr)
        return 1

    all_paragraphs = [{"index": i, "style": p.style.name if p.style is not None else None, "text": p.text} for i, p in enumerate(doc.paragraphs)]
    total_paragraphs = len(all_paragraphs)
    offset = max(args.offset, 0)
    selected = all_paragraphs[offset : offset + args.limit]

    tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables]

    # paragraphsとtablesは別々のフラットリストで返るため、両者の文書内での
    # 前後関係（どの段落とどの表が隣接しているか）が分からない。表の直前/直後に
    # 段落を挿入したい等、位置関係が必要な編集の判断材料として、本文（body）
    # 直下の子要素を文書順に並べたインデックスを別途返す。
    from docx.oxml.ns import qn

    p_index_by_elm = {p._p: i for i, p in enumerate(doc.paragraphs)}
    t_index_by_elm = {t._tbl: i for i, t in enumerate(doc.tables)}
    body_order = []
    for child in doc.element.body:
        if child.tag == qn("w:p") and child in p_index_by_elm:
            body_order.append({"type": "paragraph", "index": p_index_by_elm[child]})
        elif child.tag == qn("w:tbl") and child in t_index_by_elm:
            body_order.append({"type": "table", "index": t_index_by_elm[child]})

    # docx-edit/scripts/_track_changes.py を import する（1-B 相互import方式）
    _DOCX_EDIT_SCRIPTS = Path(__file__).resolve().parent.parent.parent / "docx-edit" / "scripts"
    if str(_DOCX_EDIT_SCRIPTS) not in sys.path:
        sys.path.append(str(_DOCX_EDIT_SCRIPTS))
    from _track_changes import count_revisions

    track_changes = count_revisions(doc)

    props = doc.core_properties
    core_properties = {
        "title": props.title or None,
        "author": props.author or None,
        "created": props.created.isoformat() if props.created else None,
        "modified": props.modified.isoformat() if props.modified else None,
    }

    # セクション（ページ設定）情報を取得
    sections_info = []
    for section in doc.sections:
        sections_info.append(
            {
                "page_width_cm": _length_cm(section.page_width),
                "page_height_cm": _length_cm(section.page_height),
                "left_margin_cm": _length_cm(section.left_margin),
                "right_margin_cm": _length_cm(section.right_margin),
                "top_margin_cm": _length_cm(section.top_margin),
                "bottom_margin_cm": _length_cm(section.bottom_margin),
            }
        )

    # 表の列幅を取得
    table_column_widths_list = []
    for table in doc.tables:
        col_widths = []
        for col in table.columns:
            width_cm = _length_cm(col.width) if col.width is not None else None
            col_widths.append(width_cm)
        table_column_widths_list.append(col_widths)

    # 挿入画像の幅を取得
    inline_image_widths_cm = []
    for inline_shape in doc.inline_shapes:
        width_cm = _length_cm(inline_shape.width) if inline_shape.width is not None else None
        inline_image_widths_cm.append(width_cm)

    # 警告を生成
    warnings = []
    if sections_info:
        first_section = sections_info[0]
        page_width_cm = first_section["page_width_cm"]
        left_margin_cm = first_section["left_margin_cm"]
        right_margin_cm = first_section["right_margin_cm"]

        if page_width_cm is not None and left_margin_cm is not None and right_margin_cm is not None:
            usable_width_cm = page_width_cm - left_margin_cm - right_margin_cm

            # 表の列幅チェック
            for table_idx, col_widths in enumerate(table_column_widths_list, 1):
                if all(w is not None for w in col_widths):
                    total_width = sum(col_widths)
                    if total_width > usable_width_cm:
                        msg = f"表{table_idx}の列幅合計({total_width:.2f}cm)が" f"利用可能なページ幅({usable_width_cm:.2f}cm)を超えています"
                        warnings.append(msg)

            # 挿入画像のサイズチェック
            for img_idx, width_cm in enumerate(inline_image_widths_cm, 1):
                if width_cm is not None and width_cm > usable_width_cm:
                    msg = f"挿入画像({img_idx}番目)の幅({width_cm:.2f}cm)が" f"ページ幅({usable_width_cm:.2f}cm)を超えています"
                    warnings.append(msg)

    result = {
        "path": str(path),
        "total_paragraphs": total_paragraphs,
        "start_index": offset if selected else None,
        "end_index": offset + len(selected) - 1 if selected else None,
        "paragraphs": selected,
        "table_count": len(tables),
        "tables": tables,
        "body_order": body_order,
        "core_properties": core_properties,
        "track_changes": track_changes,
        "sections": sections_info,
        "table_column_widths": table_column_widths_list,
        "inline_image_widths_cm": inline_image_widths_cm,
    }
    if warnings:
        result["warnings"] = warnings
    summary = summarize_result(result, ["paragraphs", "tables", "body_order"])
    summary.update(write_json_result(result, "docx_read", path))
    print(json.dumps(summary, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
