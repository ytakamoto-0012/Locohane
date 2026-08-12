"""edit_docx.py が適用する「操作（op）」のディスパッチ実装。

各関数は (DocEditContext, op_dict) を受け取り、python-docx の Document へ
副作用を適用する。戻り値は呼び出し元へ報告する結果（dict）または None。
DocEditContextは段落indexを「このopsバッチ開始時点の番号」として解決し、
delete_paragraphによる段落数の変化を自動追跡する（詳細はクラス docstring 参照）。

run_script からは直接実行されない。edit_docx.py から import して使う。
"""

from __future__ import annotations

from _blocks import (
    BLOCK_HANDLERS,
    HEADING_FONT,
    _apply_run_props,
    _set_cell_shading,
    _set_east_asian_font,
    _set_paragraph_left_border,
    _set_paragraph_shading,
    resolve_theme,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH

from _track_changes import (
    DEFAULT_AUTHOR,
    accept_all_changes,
    apply_tracked_replace_to_run,
    is_plain_text_run,
    now_iso,
    reject_all_changes,
)


class DocEditContext:
    """opsバッチ全体で共有する状態。

    `index`/`paragraph_index`は常に「このopsバッチを開始した時点（＝直前の
    docx-readが見せていたはずの状態）の段落番号」として解決する。
    `delete_paragraph`で段落数が変わっても、それより後のopで毎回シフト量を
    手計算する必要が無い（pptx-editのスライド番号解決と同じ考え方。
    Excel編集での行挿入後の絶対行番号ズレ事故と同根の問題への対応）。
    """

    def __init__(self, doc):
        self.doc = doc
        # 段落のXML要素(<w:p>)は削除されない限り同一オブジェクトのままなので、
        # Python識別性でのlist.index()による生存追跡が可能。
        self._original_paragraph_elements = [p._p for p in doc.paragraphs]

    def get_paragraph(self, index: int):
        idx = int(index)
        total = len(self._original_paragraph_elements)
        if not (0 <= idx < total):
            raise ValueError(f"存在しないindexです: {index}（このopsバッチ開始時点の総段落数: {total}）")
        element = self._original_paragraph_elements[idx]
        live_elements = [p._p for p in self.doc.paragraphs]
        try:
            live_idx = live_elements.index(element)
        except ValueError:
            raise ValueError(
                f"段落{index}は、この呼び出し内の先行する操作（delete_paragraph等）で既に削除されています"
            )
        return self.doc.paragraphs[live_idx]


def _apply_plain_replace_to_run(run, old_text: str, new_text: str, occurrence: str) -> int:
    text = run.text
    if old_text not in text:
        return 0
    count = 1 if occurrence == "first" else text.count(old_text)
    run.text = text.replace(old_text, new_text, 1 if occurrence == "first" else -1)
    return count


def op_find_replace(ctx: DocEditContext, op: dict) -> dict:
    doc = ctx.doc
    old_text = op["old_text"]
    if not old_text:
        raise ValueError("old_text は空文字列にできません")
    new_text = op.get("new_text", "")
    track = bool(op.get("track_changes", False))
    occurrence = op.get("occurrence", "all")
    if occurrence not in ("all", "first"):
        raise ValueError(f"occurrence は 'all' または 'first' である必要があります: {occurrence!r}")
    author = op.get("author") or DEFAULT_AUTHOR
    date = op.get("date") or now_iso()

    if op.get("paragraph_index") is not None:
        targets = [ctx.get_paragraph(op["paragraph_index"])]
    else:
        targets = doc.paragraphs

    replaced = 0
    for para in targets:
        for run in list(para.runs):
            if not is_plain_text_run(run._r):
                continue
            if track:
                n = apply_tracked_replace_to_run(doc, run._r, old_text, new_text, occurrence, author, date)
            else:
                n = _apply_plain_replace_to_run(run, old_text, new_text, occurrence)
            replaced += n
            if occurrence == "first" and replaced:
                break
        if occurrence == "first" and replaced:
            break

    if replaced == 0:
        raise ValueError(
            f"old_text が見つかりませんでした: {old_text!r}"
            "（1つのrun内に収まっている必要があります。書式が混在する文字列や複数runにまたがる文字列、"
            "既存のTrack Changes内の文字列は検出できません）"
        )
    return {"op": "find_replace", "replaced_count": replaced}


def op_append_block(ctx: DocEditContext, op: dict) -> dict:
    block = op["block"]
    if not isinstance(block, dict):
        raise ValueError("block はオブジェクトである必要があります")
    handler = BLOCK_HANDLERS.get(block.get("type"))
    if handler is None:
        raise ValueError(f"未対応のblock typeです: {block.get('type')!r}（対応: {sorted(BLOCK_HANDLERS)}）")
    theme = resolve_theme(op.get("theme"))
    handler(ctx.doc, block, theme)
    return {"op": "append_block", "block_type": block.get("type")}


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def op_set_paragraph_style(ctx: DocEditContext, op: dict) -> dict:
    """既存段落を明示的に再配色・再配置する（ユーザーが見た目の変更を明示的に頼んだ場合のみ使う）。

    `role`（"heading"/"callout"）+ `theme` で定番の見た目を指定するか、
    `color`/`bold`/`italic`/`underline`/`font`/`size_pt`（テキスト）・
    `fill_color`/`border_color`（段落の背景・左罫線）・`align`（left/center/
    right/justify）・`indent_left_cm`（左インデント）・`line_spacing`
    （行間の倍率、1.0=シングル・1.5・2.0等）を個別に指定する。
    明示キーはroleから導いた既定値より優先される。
    """
    index = op["index"]
    idx = int(index)
    para = ctx.get_paragraph(index)

    style = dict(op)
    role = op.get("role")
    if role is not None:
        theme = resolve_theme(op.get("theme"))
        if role == "heading":
            style.setdefault("color", theme["primary"])
            style.setdefault("bold", True)
            style.setdefault("font", HEADING_FONT)
        elif role == "callout":
            style.setdefault("fill_color", theme["secondary"])
            style.setdefault("border_color", theme["primary"])
        else:
            raise ValueError(f"未対応の role です: {role!r}（対応: heading, callout）")

    align = style.get("align")
    if align is not None and align not in _ALIGN_MAP:
        raise ValueError(f"未対応の align です: {align!r}（対応: {', '.join(_ALIGN_MAP)}）")

    has_any = any(
        style.get(k) is not None
        for k in ("color", "bold", "italic", "underline", "font", "size_pt",
                   "fill_color", "border_color", "align", "indent_left_cm", "line_spacing")
    )
    if not has_any:
        raise ValueError(
            "color/bold/italic/underline/font/size_pt/fill_color/border_color/"
            "align/indent_left_cm/line_spacing のいずれか、"
            "または role（heading/callout）+ theme の指定が必要です"
        )

    if style.get("fill_color"):
        _set_paragraph_shading(para, str(style["fill_color"]).lstrip("#").upper())
    if style.get("border_color"):
        _set_paragraph_left_border(para, str(style["border_color"]).lstrip("#").upper())
    if align is not None:
        para.alignment = _ALIGN_MAP[align]
    if style.get("indent_left_cm") is not None:
        from docx.shared import Cm
        para.paragraph_format.left_indent = Cm(float(style["indent_left_cm"]))
    if style.get("line_spacing") is not None:
        para.paragraph_format.line_spacing = float(style["line_spacing"])
    for run in para.runs:
        _apply_run_props(run, style)

    return {"op": "set_paragraph_style", "index": idx}


def op_set_table_style(ctx: DocEditContext, op: dict) -> dict:
    """既存の表の行を明示的に再配色する（ユーザーが見た目の変更を明示的に頼んだ場合のみ使う）。

    既定では見出し行（0行目）のみを対象にする。`row`（0始まり行番号）または
    `all_rows`:true を指定すると対象を変えられる。見出し行は太字＋見出し用
    フォントも当てるが、`row`/`all_rows`指定時の本体行は塗り・文字色のみ
    （太字強制はしない）。
    """
    from docx.shared import RGBColor

    table_index = op["table_index"]
    tables = ctx.doc.tables
    idx = int(table_index)
    if not (0 <= idx < len(tables)):
        raise ValueError(f"存在しないtable_indexです: {table_index}（総表数: {len(tables)}）")
    table = tables[idx]

    theme = resolve_theme(op["theme"]) if op.get("theme") else None
    fill_color = op.get("header_fill") or (theme["primary"] if theme else None)
    font_color = op.get("header_font_color") or (theme["text_on_primary"] if theme else None)
    if not fill_color:
        raise ValueError("set_table_style には theme または header_fill のいずれかが必要です")
    fill_color = str(fill_color).lstrip("#").upper()
    font_color = str(font_color).lstrip("#").upper() if font_color else None

    if op.get("all_rows"):
        target_rows = range(len(table.rows))
        force_bold = False
    elif op.get("row") is not None:
        row = int(op["row"])
        if not (0 <= row < len(table.rows)):
            raise ValueError(f"存在しない行です: {row}（この表は{len(table.rows)}行）")
        target_rows = [row]
        force_bold = False
    else:
        target_rows = [0]
        force_bold = True

    for r in target_rows:
        for cell in table.rows[r].cells:
            _set_cell_shading(cell, fill_color)
            for para in cell.paragraphs:
                for run in para.runs:
                    if force_bold:
                        run.bold = True
                    if font_color:
                        run.font.color.rgb = RGBColor.from_string(font_color)
                    if force_bold:
                        _set_east_asian_font(run.font, run._element, HEADING_FONT)

    return {"op": "set_table_style", "table_index": idx}


def op_delete_paragraph(ctx: DocEditContext, op: dict) -> dict:
    index = op["index"]
    if op.get("track_changes"):
        raise ValueError(
            "delete_paragraph はtrack_changesに対応していません（段落マーク削除はOOXML上複雑なため非対応です）。"
            "変更履歴として残したい場合は、段落内の全文をfind_replaceでtrack_changes:trueにより"
            "空文字へ置換してください（段落自体は残り、中身の削除のみが変更履歴になります）"
        )
    para = ctx.get_paragraph(index)
    p_element = para._p
    p_element.getparent().remove(p_element)
    return {"op": "delete_paragraph", "deleted_index": int(index)}


def op_accept_all_changes(ctx: DocEditContext, op: dict) -> dict:
    result = accept_all_changes(ctx.doc)
    return {"op": "accept_all_changes", **result}


def op_reject_all_changes(ctx: DocEditContext, op: dict) -> dict:
    result = reject_all_changes(ctx.doc)
    return {"op": "reject_all_changes", **result}


OP_HANDLERS = {
    "find_replace": op_find_replace,
    "append_block": op_append_block,
    "set_paragraph_style": op_set_paragraph_style,
    "set_table_style": op_set_table_style,
    "delete_paragraph": op_delete_paragraph,
    "accept_all_changes": op_accept_all_changes,
    "reject_all_changes": op_reject_all_changes,
}


def apply_op(ctx: DocEditContext, op: dict) -> dict | None:
    op_name = op.get("op")
    handler = OP_HANDLERS.get(op_name)
    if handler is None:
        raise ValueError(f"未対応のopです: {op_name!r}（対応op: {sorted(OP_HANDLERS)}）")
    return handler(ctx, op)
