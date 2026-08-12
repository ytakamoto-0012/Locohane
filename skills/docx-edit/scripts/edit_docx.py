"""ops（操作）のリストを適用して既存のdocxファイルを編集する。

docx-tools スキルの実行スクリプト（progressive disclosure 第3段階）。
run_script から
    python edit_docx.py <docx_path> --ops-json "<JSON配列>" [--output <別名保存先>]
または
    python edit_docx.py <docx_path> --ops-file <JSONファイルのパス> [--output ...]
の形で呼ばれる（このプロジェクトには汎用のファイル書き込みツールが無いため、
LLMが組み立てたJSONをそのまま --ops-json 引数として渡せるようにしている。
excel-tools の edit_excel.py と同じ設計）。

このスクリプトは既存ファイルの編集専用（--newは持たない）。新規作成は
create_docx.py を使うこと。--output を省略した場合は <docx_path> へ
上書き保存する。

ops の各要素の形式や対応opの一覧は SKILL.md を参照。実装（各opの処理）は
scripts/_ops.py の OP_HANDLERS を参照。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from _common import backup_before_overwrite, register_output_path, setup_utf8_stdio
from _ops import DocEditContext, apply_op


def _load_json_arg(args: argparse.Namespace) -> str:
    if args.ops_json is not None:
        return args.ops_json
    data_path = Path(args.ops_file)
    if not data_path.is_file():
        raise FileNotFoundError(f"opsファイルが見つかりません: {args.ops_file}")
    try:
        return data_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return data_path.read_text(encoding="cp932")


def main() -> int:
    setup_utf8_stdio()
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path", help="編集対象の既存.docxファイルのパス")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--ops-json", help="適用する操作のJSON配列をそのまま渡す")
    group.add_argument("--ops-file", help="適用する操作のJSON配列を書いたファイルのパス")
    parser.add_argument("--output", default=None, help="保存先パス（省略時は docx_path へ上書き保存）")
    args = parser.parse_args()

    path = Path(args.docx_path)
    if not path.exists():
        print(f"編集対象のファイルが見つかりません: {args.docx_path}", file=sys.stderr)
        return 1
    if path.is_dir():
        print(f"指定パスはディレクトリです（ファイル専用）: {args.docx_path}", file=sys.stderr)
        return 1
    if path.suffix.lower() == ".doc":
        print(
            "拡張子が .doc（レガシーのバイナリ形式）です。このスキルは .docx のみ対応しています。"
            "Microsoft Wordで開き「名前を付けて保存」で .docx 形式に変換してから再度お試しください。",
            file=sys.stderr,
        )
        return 1
    if path.suffix.lower() != ".docx":
        print(f"拡張子が .docx ではありません: {args.docx_path}", file=sys.stderr)
        return 1

    output_path = Path(args.output) if args.output else path

    try:
        raw = _load_json_arg(args)
    except FileNotFoundError as e:
        print(str(e), file=sys.stderr)
        return 1

    try:
        ops = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"opsのJSON解析に失敗しました: {e}", file=sys.stderr)
        return 1
    if not isinstance(ops, list):
        print("opsはJSON配列である必要があります", file=sys.stderr)
        return 1

    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        print(f".docxとして読み込めませんでした（壊れているか非対応形式の可能性）: {args.docx_path}", file=sys.stderr)
        return 1
    except OSError as e:
        print(f"ファイル読み込みに失敗しました: {e}", file=sys.stderr)
        return 1

    ctx = DocEditContext(doc)
    op_results = []
    for idx, op in enumerate(ops):
        if not isinstance(op, dict) or "op" not in op:
            print(f"ops[{idx}] が不正です（'op'キーを持つオブジェクトである必要があります）: {op!r}", file=sys.stderr)
            return 1
        try:
            result = apply_op(ctx, op)
        except (KeyError, ValueError, TypeError) as e:
            print(f"ops[{idx}]（op={op.get('op')!r}）の適用に失敗しました: {e}", file=sys.stderr)
            return 1
        op_results.append(result or {"op": op.get("op")})

    output_path.parent.mkdir(parents=True, exist_ok=True)
    backup_path = backup_before_overwrite(output_path)
    try:
        doc.save(str(output_path))
    except OSError as e:
        print(f"ファイルの保存に失敗しました: {e}", file=sys.stderr)
        return 1

    result = {
        "path": str(output_path),
        "backup_path": str(backup_path) if backup_path else None,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "applied_ops": len(ops),
        "op_results": op_results,
    }
    path_memory = register_output_path(output_path, description="edit_docxが生成/更新したDOCX")
    if path_memory:
        result["path_memory"] = path_memory
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
