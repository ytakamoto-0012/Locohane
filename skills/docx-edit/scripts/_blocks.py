"""docxの「ブロック」（見出し・段落・箇条書き・表・画像・改ページ）を組み立てる共通ロジック。

create_docx.py（新規作成）と _ops.py（既存ファイル編集の append_block op）の
両方から import される。ブロックのJSON仕様は SKILL.md を参照。

run_script からは直接実行されない。
"""

from __future__ import annotations

import sys
from pathlib import Path

DEFAULT_FONT = "游明朝"
HEADING_FONT = "游ゴシック"

# _shared/office_theme.py から THEMES / resolve_theme を import する（1-B 相互import方式）
_OFFICE_SHARED = Path(__file__).resolve().parent.parent.parent / "_shared"
if str(_OFFICE_SHARED) not in sys.path:
    sys.path.append(str(_OFFICE_SHARED))
from office_theme import DEFAULT_THEME, THEMES, resolve_theme  # noqa: E402


def _set_cell_shading(cell, color_hex: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_paragraph_shading(paragraph, color_hex: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def _set_paragraph_left_border(paragraph, color_hex: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_east_asian_font(font_obj, element, font_name: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    font_obj.name = font_name
    rPr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element
    r_fonts = rPr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rPr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _apply_run_props(run, run_spec: dict) -> None:
    # keyの存在とNone判定を明示（False/0 の明示指定でも正しく適用する）
    for bool_key in ("bold", "italic", "underline"):
        if bool_key in run_spec and run_spec[bool_key] is not None:
            setattr(run, bool_key, bool(run_spec[bool_key]))

    from docx.shared import Pt, RGBColor

    if "size_pt" in run_spec and run_spec["size_pt"] is not None:
        run.font.size = Pt(float(run_spec["size_pt"]))
    if run_spec.get("color"):
        run.font.color.rgb = RGBColor.from_string(str(run_spec["color"]).lstrip("#").upper())
    if run_spec.get("font"):
        _set_east_asian_font(run.font, run._element, str(run_spec["font"]))


def _add_heading(doc, block: dict, theme: dict):
    # 既存文書へのappend_blockでは、その文書自身のWordテーマ（組み込み
    # 見出しスタイルが参照する色）をそのまま尊重する。theme引数はここでは
    # 使わない（callout等、文書側に対応物が無い新規要素にのみ使う）。
    level = max(0, min(int(block.get("level", 1)), 9))
    doc.add_heading(str(block.get("text", "")), level=level)


def _add_paragraph(doc, block: dict, theme: dict):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para = doc.add_paragraph()
    if block.get("alignment") in align_map:
        para.alignment = align_map[block["alignment"]]

    runs = block.get("runs")
    if runs:
        for run_spec in runs:
            run = para.add_run(str(run_spec.get("text", "")))
            _apply_run_props(run, run_spec)
    else:
        run = para.add_run(str(block.get("text", "")))
        _apply_run_props(run, block)


def _add_list(doc, block: dict, theme: dict, style_name: str):
    for item in block.get("items") or []:
        doc.add_paragraph(str(item), style=style_name)


def _add_table(doc, block: dict, theme: dict):
    # heading同様、theme引数は使わず太字のみ（既存文書に配色を押し付けない）。
    # 後から配色したい場合は set_table_style op を明示的に呼ぶ。
    rows = block.get("rows") or []
    if not rows:
        raise ValueError("table.rows が空です（少なくとも1行必要）")
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    header_row = bool(block.get("header_row"))
    for i, row in enumerate(rows):
        for j in range(n_cols):
            value = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = "" if value is None else str(value)
            if header_row and i == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _add_image(doc, block: dict, theme: dict):
    from docx.shared import Cm

    image_path = Path(str(block.get("path", "")))
    if not image_path.is_file():
        raise ValueError(f"画像ファイルが見つかりません: {block.get('path')}")
    kwargs = {}
    if block.get("width_cm"):
        kwargs["width"] = Cm(float(block["width_cm"]))
    if block.get("height_cm"):
        kwargs["height"] = Cm(float(block["height_cm"]))
    doc.add_picture(str(image_path), **kwargs)


def _add_callout(doc, block: dict, theme: dict):
    """左に太いアクセント罫線＋薄い塗りの強調ボックス段落（docx-createのcalloutと同じ）。"""
    para = doc.add_paragraph()
    _set_paragraph_shading(para, theme["secondary"])
    _set_paragraph_left_border(para, theme["primary"])
    run = para.add_run(str(block.get("text", "")))
    _apply_run_props(run, block)


BLOCK_HANDLERS = {
    "heading": _add_heading,
    "paragraph": _add_paragraph,
    "bullet_list": lambda doc, b, theme: _add_list(doc, b, theme, "List Bullet"),
    "number_list": lambda doc, b, theme: _add_list(doc, b, theme, "List Number"),
    "table": _add_table,
    "image": _add_image,
    "callout": _add_callout,
    "page_break": lambda doc, b, theme: doc.add_page_break(),
}
