"""docxファイルを読み込み、段落・表・文書プロパティを JSON で出力する。

docx-tools スキルの実行スクリプト（progressive disclosure 第3段階）。
run_script から
    python read_docx.py <docx_path> [--offset N] [--limit N]
の形で呼ばれる。

read_file.py の offset/limit と同じ考え方を段落単位に適用し、
1回の呼び出しで返す段落数の上限を設けて巨大な出力を避ける
（表は全件返す。読み込みのみで書き込みは行わないため python-docx の
「既存ファイルへの変更」という性質は関係ない）。

.doc（レガシーのバイナリ形式）は python-docx が非対応のため扱えない。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from _common import setup_utf8_stdio


def main() -> int:
    setup_utf8_stdio()
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path")
    parser.add_argument("--offset", type=int, default=0)
    parser.add_argument("--limit", type=int, default=300)
    args = parser.parse_args()

    path = Path(args.docx_path)
    if not path.exists():
        print(f"ファイルが見つかりません: {args.docx_path}", file=sys.stderr)
        return 1
    if path.is_dir():
        print(f"指定パスはディレクトリです（ファイル専用）: {args.docx_path}", file=sys.stderr)
        return 1
    if path.suffix.lower() == ".doc":
        print(
            "拡張子が .doc（レガシーのバイナリ形式）です。このスキルは .docx のみ対応しています。"
            "Microsoft Wordで開き「名前を付けて保存」で .docx 形式に変換してから再度お試しください。",
            file=sys.stderr,
        )
        return 1
    if path.suffix.lower() != ".docx":
        print(f"拡張子が .docx ではありません: {args.docx_path}", file=sys.stderr)
        return 1

    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        print(f".docxとして読み込めませんでした（壊れているか非対応形式の可能性）: {args.docx_path}", file=sys.stderr)
        return 1
    except OSError as e:
        print(f"ファイル読み込みに失敗しました: {e}", file=sys.stderr)
        return 1

    all_paragraphs = [
        {"index": i, "style": p.style.name if p.style is not None else None, "text": p.text}
        for i, p in enumerate(doc.paragraphs)
    ]
    total_paragraphs = len(all_paragraphs)
    offset = max(args.offset, 0)
    selected = all_paragraphs[offset : offset + args.limit]

    tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables]

    from _track_changes import count_revisions

    track_changes = count_revisions(doc)

    props = doc.core_properties
    core_properties = {
        "title": props.title or None,
        "author": props.author or None,
        "created": props.created.isoformat() if props.created else None,
        "modified": props.modified.isoformat() if props.modified else None,
    }

    result = {
        "path": str(path),
        "total_paragraphs": total_paragraphs,
        "start_index": offset if selected else None,
        "end_index": offset + len(selected) - 1 if selected else None,
        "paragraphs": selected,
        "table_count": len(tables),
        "tables": tables,
        "core_properties": core_properties,
        "track_changes": track_changes,
    }
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
