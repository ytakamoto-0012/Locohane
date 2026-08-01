"""JSON仕様からdocxファイルを新規生成する。

docx-tools スキルの実行スクリプト（progressive disclosure 第3段階）。
run_script から
    python create_docx.py <output_path> --data "<JSON文字列>"
または
    python create_docx.py <output_path> --data-file <JSONファイルのパス>
の形で呼ばれる（このプロジェクトには汎用のファイル書き込みツールが無いため、
LLMが組み立てたJSONをそのまま --data 引数として渡せるようにしている。
pptx-tools の create_pptx.py と同じ設計。出力先が既に存在する場合は
pdf-tools/pptx-tools と同様に常に上書きする）。

JSON仕様の形式は SKILL.md を参照。ページ番号フィールド（PAGE）は
python-docx に高レベルAPIが無いため oxml を直接組み立てて挿入する。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from _common import setup_utf8_stdio
from _blocks import BLOCK_HANDLERS, DEFAULT_FONT, _set_east_asian_font

PAGE_SIZES_CM = {
    "a4": (21.0, 29.7),
    "letter": (21.59, 27.94),
}


def _apply_page(doc, page_spec: dict):
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm

    section = doc.sections[0]
    size_key = str(page_spec.get("size", "a4")).lower()
    width_cm, height_cm = PAGE_SIZES_CM.get(size_key, PAGE_SIZES_CM["a4"])
    if str(page_spec.get("orientation", "portrait")).lower() == "landscape":
        width_cm, height_cm = height_cm, width_cm
        section.orientation = WD_ORIENT.LANDSCAPE

    section.page_width = Cm(width_cm)
    section.page_height = Cm(height_cm)

    margin = page_spec.get("margin_cm") or {}
    section.top_margin = Cm(float(margin.get("top", 2.5)))
    section.bottom_margin = Cm(float(margin.get("bottom", 2.5)))
    section.left_margin = Cm(float(margin.get("left", 3.0)))
    section.right_margin = Cm(float(margin.get("right", 2.5)))


def _add_page_number_field(paragraph) -> None:
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _apply_header_footer(doc, spec: dict):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]
    header_text = spec.get("header_text")
    if header_text:
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = str(header_text)

    footer_text = spec.get("footer_text")
    page_number = bool(spec.get("page_number"))
    if footer_text or page_number:
        section.footer.is_linked_to_previous = False
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_text:
            para.add_run(str(footer_text) + ("　" if page_number else ""))
        if page_number:
            _add_page_number_field(para)


def _build_document(spec: dict):
    from docx import Document

    doc = Document()

    # 既定フォントを東アジア用に設定しないと日本語が意図しないフォントで
    # 表示されることがあるため、Normalスタイルに明示しておく。
    normal_style = doc.styles["Normal"]
    _set_east_asian_font(normal_style.font, normal_style.element.get_or_add_rPr(), DEFAULT_FONT)

    if spec.get("page"):
        _apply_page(doc, spec["page"])
    _apply_header_footer(doc, spec)

    core_props = spec.get("core_properties") or {}
    if core_props.get("title"):
        doc.core_properties.title = str(core_props["title"])
    if core_props.get("author"):
        doc.core_properties.author = str(core_props["author"])

    warnings: list[str] = []
    for i, block in enumerate(spec.get("blocks") or []):
        block_type = block.get("type") if isinstance(block, dict) else None
        handler = BLOCK_HANDLERS.get(block_type)
        if handler is None:
            warnings.append(f"blocks[{i}]: 未知の type '{block_type}' はスキップしました")
            continue
        handler(doc, block)

    return doc, warnings


def main() -> int:
    setup_utf8_stdio()
    parser = argparse.ArgumentParser()
    parser.add_argument("output_path")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--data", help="文書仕様のJSON文字列をそのまま渡す")
    group.add_argument("--data-file", help="文書仕様を書いたJSONファイルのパス")
    args = parser.parse_args()

    output_path = Path(args.output_path)
    if output_path.suffix.lower() != ".docx":
        print(f"出力パスの拡張子は .docx にしてください: {args.output_path}", file=sys.stderr)
        return 1

    if args.data is not None:
        raw = args.data
    else:
        data_path = Path(args.data_file)
        if not data_path.is_file():
            print(f"データ仕様ファイルが見つかりません: {args.data_file}", file=sys.stderr)
            return 1
        try:
            raw = data_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw = data_path.read_text(encoding="cp932")

    try:
        spec = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"データ仕様JSONの解析に失敗しました: {e}", file=sys.stderr)
        return 1

    try:
        doc, warnings = _build_document(spec)
    except ValueError as e:
        print(str(e), file=sys.stderr)
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(output_path))
    except OSError as e:
        print(f"ファイルの保存に失敗しました: {e}", file=sys.stderr)
        return 1

    result = {
        "output_path": str(output_path),
        "blocks_written": len(spec.get("blocks") or []),
        "warnings": warnings,
    }
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
